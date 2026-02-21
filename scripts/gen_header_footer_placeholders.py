"""
Generate a .docx with placeholders in header and footer.

Tests that compile-time extraction reads section_property headers/footers.
"""

from docx import Document
from docx.shared import Pt

doc = Document()

header = doc.sections[0].header
header.is_linked_to_previous = False
header.paragraphs[0].text = "Header: {header_title}"

footer = doc.sections[0].footer
footer.is_linked_to_previous = False
footer.paragraphs[0].text = "Footer: {footer_page}"

doc.add_paragraph("Body content only")

out_path = "test-crate/templates/header_footer_placeholders.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
