"""
Generate a .docx where a placeholder is split across multiple <w:r> runs.

Word commonly splits text like {FirstName} into separate runs due to
spell-checking, language detection, or formatting changes. For example:
  Run 1: "{First"
  Run 2: "Name}"

This is the #1 edge case for docx template engines (per OOXML spec §17.3.2).
"""

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import lxml.etree as ET

doc = Document()
p = doc.add_paragraph()

# Manually build runs that split the placeholder across boundaries.
# We want the paragraph to logically read:
#   "Hello {FirstName}, welcome to {CompanyName}!"
# But we split {FirstName} into 3 runs and {CompanyName} into 2 runs.

fragments = [
    ("Hello ", None),
    ("{First", None),       # placeholder start
    ("Name}", None),        # placeholder end
    (", welcome to ", None),
    ("{Company", None),     # placeholder start
    ("Name}!", None),       # placeholder end — note the "!" is glued to the closing brace
]

for text, _ in fragments:
    run = p.add_run(text)

# Verify the raw XML looks right
print("Generated XML for paragraph:")
print(ET.tostring(p._element, pretty_print=True).decode())

out_path = "test-crate/templates/split_runs_template.docx"
doc.save(out_path)
print(f"\nSaved to {out_path}")
