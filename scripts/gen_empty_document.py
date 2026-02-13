"""
Generate edge-case .docx files with minimal or no content.

Tests:
- Completely empty document (no paragraphs added)
- Document with empty paragraphs only
- Document with whitespace-only text
- Document with curly braces but no valid placeholder pattern

These should all result in structs with zero fields (unit structs).
"""

from docx import Document

# 1. Completely empty document — python-docx always includes one empty paragraph
doc1 = Document()
doc1.save("test-crate/templates/empty_document.docx")
print("Saved empty_document.docx")

# 2. Document with only blank paragraphs
doc2 = Document()
doc2.add_paragraph("")
doc2.add_paragraph("")
doc2.add_paragraph("")
doc2.save("test-crate/templates/blank_paragraphs.docx")
print("Saved blank_paragraphs.docx")

# 3. Document with braces but no valid placeholders
doc3 = Document()
doc3.add_paragraph("This has {  } empty braces.")
doc3.add_paragraph("And some { } with just a space.")
doc3.add_paragraph("Also lone { and lone } without pairs.")
doc3.add_paragraph("Nested {{double}} braces.")
doc3.save("test-crate/templates/invalid_placeholders.docx")
print("Saved invalid_placeholders.docx")
