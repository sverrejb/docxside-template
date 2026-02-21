"""
Generate a .docx with placeholders across all document areas:
body paragraphs, tables, headers, and footers.

Tests that all areas are extracted and deduplicated into a single struct.
"""

from docx import Document

doc = Document()

# Header
header = doc.sections[0].header
header.is_linked_to_previous = False
header.paragraphs[0].text = "Doc: {doc_title}"

# Footer
footer = doc.sections[0].footer
footer.is_linked_to_previous = False
footer.paragraphs[0].text = "Page {page_num}"

# Body paragraph
doc.add_paragraph("Hello {body_name}")

# Table
table = doc.add_table(rows=1, cols=2)
table.cell(0, 0).text = "{cell_label}"
table.cell(0, 1).text = "{cell_value}"

out_path = "test-crate/templates/combined_areas.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
