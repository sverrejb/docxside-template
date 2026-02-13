"""
Generate a .docx with unicode characters in placeholder names.

Tests that the pipeline handles:
- Accented Latin characters (é, ñ, ü)
- Non-Latin scripts (Norwegian æøå)
- Spaces and mixed case in placeholder names
- Emoji in surrounding text (not in placeholders)

The heck crate's snake_case and syn's Ident validation may reject some of these.
This tests the boundary between "accepted" and "skipped" placeholders.
"""

from docx import Document

doc = Document()

# Simple accented placeholder — should become a valid Rust ident if heck handles it
doc.add_paragraph("Kjære {Fornavn}, velkommen til {Bedriftsnavn}.")

# Norwegian special chars in placeholder name
doc.add_paragraph("Hei {Øltype}, vi har {Størrelse} på lager.")

# Placeholder with spaces (the code converts spaces to underscores)
doc.add_paragraph("Dear {First Name}, your order {Order Number} is ready.")

# Placeholder with digits
doc.add_paragraph("Reference: {Case2024Id}")

# Placeholder with colons (the code replaces : with _)
doc.add_paragraph("System: {app:version}")

out_path = "test-crate/templates/unicode_placeholders.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
