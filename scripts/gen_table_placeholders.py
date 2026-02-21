"""
Generate a .docx with placeholders inside table cells.

Tests that compile-time placeholder extraction traverses into
Table -> TableRow -> TableCell -> Paragraph structures.
"""

from docx import Document

doc = Document()

doc.add_paragraph("Table template")

table = doc.add_table(rows=3, cols=2)
table.cell(0, 0).text = "Label"
table.cell(0, 1).text = "Value"
table.cell(1, 0).text = "Name"
table.cell(1, 1).text = "{table_name}"
table.cell(2, 0).text = "City"
table.cell(2, 1).text = "{table_city}"

out_path = "test-crate/templates/table_placeholders.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
