"""
Generate a .docx with placeholders inside deeply nested formatting.

Word wraps text in runs with rich formatting properties (<w:rPr>).
This tests that placeholder detection works when runs have:
- Bold/italic formatting applied mid-placeholder
- Color changes within a placeholder
- Multiple formatting properties stacked

Per OOXML spec §17.3.2, each formatting change creates a new run boundary.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

doc = Document()

# Case 1: Placeholder entirely inside a bold run
p1 = doc.add_paragraph()
p1.add_run("Bold placeholder: ")
bold_run = p1.add_run("{BoldField}")
bold_run.bold = True

# Case 2: Placeholder split across bold and non-bold runs
# Simulates user bolding part of a placeholder by accident
p2 = doc.add_paragraph()
p2.add_run("Mixed formatting: ")
r1 = p2.add_run("{Mixed")
r1.bold = True
r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
r2 = p2.add_run("Format}")
r2.italic = True

# Case 3: Placeholder with underline + font size changes
p3 = doc.add_paragraph()
p3.add_run("Styled: ")
r3a = p3.add_run("{Under")
r3a.underline = WD_UNDERLINE.SINGLE
r3a.font.size = Pt(14)
r3b = p3.add_run("lined")
r3b.font.size = Pt(10)
r3c = p3.add_run("Field}")

# Case 4: Placeholder split character-by-character (worst case)
p4 = doc.add_paragraph()
p4.add_run("Char split: ")
for ch in "{ABC}":
    r = p4.add_run(ch)
    r.bold = (ch in "{}")

out_path = "test-crate/templates/formatted_runs.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
